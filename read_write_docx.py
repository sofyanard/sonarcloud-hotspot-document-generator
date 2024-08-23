from docx import Document
import re
from docx.shared import Inches
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import RGBColor
from docx.shared import Pt

def add_header(data, doc, level: int):
    try:
        # try:
        #     create_heading_style(doc, f'Heading {level}', 1)
        # except Exception as err:
        #     pass
        heading = doc.add_heading(data, level=level)
        
        # font_size: float = 10.5
        # if level == 1:
        #     font_size = 20
        # elif level == 2:
        #     font_size = 18
        # elif level == 3:
        #     font_size = 16
            
        # for run in heading.runs:
        #     set_font(run, 'Microsoft Sans Serif', font_size=font_size)
        # doc.save(filename)
        # print(f'Data successfully appended to {filename}')
        return True
    except Exception as err:
        print(f'Error adding header {level}: {err}')
        return False

def insert_to_doc(data, doc, iteration, target):
    print(f'insert_to_doc {data["key"]}')
    ret = True
    class_name= None
    
    # set cell color
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), "C8C8C8"))
    try:
        if "line" in data:
            class_name= data["component"] + " baris ke " + str(data["line"])
        else:
            class_name= data["component"]
            
        # Convert the data to a formatted string and add to the document
        add_header(class_name, doc, 4)
        
        try:
            table = doc.add_table(rows=8, cols=2)
            table.style = 'Table Grid'
            table.autofit = False 
            table.allow_autofit = False
        except Exception as err:
            print(f'Error adding table: {err}')
            ret = False
        
        try:
            table.rows[0].cells[0].text = "No"
            table.rows[0].cells[1].text = str(iteration)
            # table.rows[0].cells[0]._element.get_or_add_tcPr().append(shading_elm)
        except Exception as err:
            print(f'Error adding No: {err}')
            ret = False
        
        try:
            table.rows[1].cells[0].text = "Class"
            table.rows[1].cells[1].text = data["securityCategory"]
            # table.rows[1].cells[0]._element.get_or_add_tcPr().append(shading_elm)
        except Exception as err:
            print(f'Error adding Class: {err}')
            ret = False
        
        try:
            table.rows[2].cells[0].text = "Severity"
            table.rows[2].cells[1].text = data["vulnerabilityProbability"]
            # table.rows[2].cells[0]._element.get_or_add_tcPr().append(shading_elm)
        except Exception as err:
            print(f'Error adding Severity: {err}')
            ret = False
        
        try:
            table.rows[3].cells[0].text = "Target"
            table.rows[3].cells[1].text = target.title()
            # table.rows[3].cells[0]._element.get_or_add_tcPr().append(shading_elm)
        except Exception as err:
            print(f'Error adding Title: {err}')
            ret = False
        
        try:
            table.rows[4].cells[0].text = "Sub Target"
            table.rows[4].cells[1].text = "Sakti Application"
            # table.rows[4].cells[0]._element.get_or_add_tcPr().append(shading_elm)
        except Exception as err:
            print(f'Error adding Title: {err}')
            ret = False

        try:
            table.rows[5].cells[0].text = "Title"
            table.rows[5].cells[1].text = data["message"]
            # table.rows[5].cells[0]._element.get_or_add_tcPr().append(shading_elm)
        except Exception as err:
            print(f'Error adding Title: {err}')
            ret = False
        
        try:
            table.rows[6].cells[0].text = "Class Name"
            table.rows[6].cells[1].text = class_name
            # table.rows[6].cells[0]._element.get_or_add_tcPr().append(shading_elm)
        except Exception as err:
            print(f'Error adding Class Name: {err}')
            ret = False
        
        try:
            table.rows[7].cells[0].text = "Detail"
            if "flows" in data and len(data["flows"]) > 0:
                i = 0
                for flow in data["flows"]:
                    msg = data["flows"][i]["locations"][0]["msg"]
                    startLine = data["flows"][i]["locations"][0]["textRange"]["startLine"]
                    endLine = data["flows"][i]["locations"][0]["textRange"]["endLine"]
                    if (startLine == endLine):
                        str_detail = f'{msg} (baris {startLine})'
                    else:
                        str_detail = f'{msg} (baris {startLine} - {endLine})'
                    insert_detail(table.rows[7].cells[1], str_detail)
                    i+=1
            # else:
            #     class_name= data["component"]
            # insert_detail(table.rows[7].cells[1], data["rule"]["htmlDesc"])
            # table.rows[7].cells[0]._element.get_or_add_tcPr().append(shading_elm)
        except Exception as err:
            print(f'Error adding Detail: {err}')
            ret = False
            
        # Set column widths (in Inches)
        column_widths = [Inches(1.5), Inches(4.0)]

        # Ensure column widths are set in the table grid
        set_column_widths_via_tblGrid(table, column_widths)
        
        # Adjust the width of each column
        for i in range(2):  # Assuming you have only 2 columns
            for cell in table.columns[i].cells:
                cell.width = column_widths[i]
                
        # change font style inside a table
        for row in table.rows:
            first_cell = row.cells[0]
            set_cell_background(first_cell, "454545")  # Grey color code
            set_cell_font_color(first_cell, RGBColor(255, 255, 255))  # White color
            
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_before = Pt(4)  # Space before paragraph
                    paragraph_format.space_after = Pt(4)   # Space after paragraph
                    paragraph_format.line_spacing = 1   # Line spacing within paragraph
                    
                    for run in paragraph.runs:
                        run.font.name = 'Microsoft Sans Serif'  # Set font name
                        run.font.size = Pt(10.5)  # Set font size

        paragraph = doc.add_paragraph()
        paragraph.add_run("\n")
    except Exception as err:
        print(f'Error appending data to Word file: {err}')
        ret = False
    return ret

def set_column_widths_via_tblGrid(table, widths):
    tbl = table._tbl
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        tblGrid.clear()  # Clear any existing gridCol elements
    else:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.insert(0, tblGrid)
    for width in widths:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(width * 1440)))  # Width in twips
        tblGrid.append(gridCol)

def initialize_doc(filename):
    try:
        doc = Document()
        # Check if the file already exists to continue appending
        try:
            doc = Document(filename)
            h1 = doc.styles['Heading 1']
            rFonts = h1.element.rPr.rFonts
            rFonts.set(qn("w:asciiTheme"), "Microsoft Sans Serif")
        except Exception:
            pass
        doc.add_heading('Sakti Fuse Main', 0)
        doc.save(filename)
        print(f'Successfully initialize {filename}')
        return doc
    except Exception as err:
        print(f'Error initialize Word file: {err}')
    
def load_doc(filename):
    try:
        doc = Document(filename)
        return doc
    except Exception as err:
        print(f'Error loading file: {repr(err)}')
    return False

def merge_and_save_docx(doc1, doc2, output_path):
    try:
        # Append the content of doc2 to doc1
        for element in doc2.element.body:
            doc1.element.body.append(element)

        doc1.save(output_path)
        return True
    except Exception as err:
        print(f'Error merge file: {repr(err)}')
    return False

def replace_text_in_docx(doc, old_text, new_text, type = "all"):
    try:
        if type == "paragraph":
            # Replace text in paragraphs
            for paragraph in doc.paragraphs:
                replace_text_in_paragraph(paragraph, old_text, new_text)
        elif type == "table":
            # Replace text in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, old_text, new_text)
        elif type == "textbox":
            replace_text_in_textboxes(doc, old_text, new_text)
        else:
            # Replace text in both paragraphs and tables
            for paragraph in doc.paragraphs:
                replace_text_in_paragraph(paragraph, old_text, new_text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, old_text, new_text)
            replace_text_in_textboxes(doc, old_text, new_text)
        return True
    except Exception as err:
        print(repr(err))
        return False

def replace_text_in_paragraph(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        inline = paragraph.runs
        for i in range(len(inline)):
            if old_text in inline[i].text:
                text = inline[i].text.replace(old_text, new_text)
                inline[i].text = text
                
def replace_text_in_textboxes(doc, old_text, new_text):
    for shape in doc.element.body.iter(qn('w:t')):
        for text in shape.iter(qn('w:t')):
            if old_text in text.text:
                text.text = text.text.replace(old_text, new_text)

def insert_detail(cell, paragraph_text):
    # Regex to find <code>...</code> and <pre>...</pre> blocks
    code_pattern = re.compile(r'(<code>.*?</code>)|(<pre.*?>.*?</pre>)', re.DOTALL)
    
    parts = None
    try:
        # Split the text based on the <code>...</code> and <pre>...</pre> blocks
        parts = code_pattern.split(paragraph_text)
    except Exception as err:
        print(f'Error Spliting Part: {err}')
    
    para = cell.add_paragraph()

    for part in parts:
        if part:  # Check if part is not None and not an empty string
            if (part.startswith('<code>') and part.endswith('</code>')) or (part.startswith('<pre') and part.endswith('</pre>')):
                try:
                    part = re.sub(r'<.*?>', '', part)
                except Exception as err:
                    print(f'Error cleanse HTML tag: {err}')
                
                try:
                    # Add the code text with grey background
                    run = para.add_run(part)
                    run.font.color.rgb = RGBColor(255, 255, 255)  # Set font color to white
                    run.font.name = 'Courier New'  # Change to any monospace font

                    # Create a shading element and set its background color to dark grey
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:val'), 'clear')
                    shading_elm.set(qn('w:color'), 'auto')
                    shading_elm.set(qn('w:fill'), 'a1a1a1')  # Dark grey color
                    run._r.get_or_add_rPr().append(shading_elm)
                except Exception as err:
                    print(f'Error adding shade: {err}')
                
            # elif (part.startswith('<strong>') and part.endswith('</strong>')):
            #     # Add the bold text
            #     # strong_content = re.sub(r'<(/)?strong.*?>', '', part)
            #     part = re.sub(r'<.*?>', '', part)
            #     run = para.add_run(part)
            #     run.bold = True
            else:
                # Remove any HTML tags from the part
                try:
                    part = re.sub(r'<.*?>', '', part)
                except Exception as err:
                    print(f'Error cleanse HTML tag: {err}')
                    
                try:
                    # Add the cleaned text
                    para.add_run(part)
                except Exception as err:
                    print(f'Error adding clean paragraph: {err}')

def set_font(run, font_name, font_size=None):
    run.font.name = font_name
    r = run._element
    rFonts = r.find(qn('w:rPr')).find(qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    if font_size:
        run.font.size = Pt(font_size)
        
def set_cell_background(cell, color):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._element.get_or_add_tcPr().append(shading_elm)
    
def set_cell_font_color(cell, color):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = color